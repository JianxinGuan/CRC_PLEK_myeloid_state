from docx import Document
from pathlib import Path

doc_path = Path(r"G:\New_CRC_Platelet\23_形成论文仓库\manuscript\CRC_PLEK_myeloid_state_manuscript_major_revision.docx")
doc = Document(doc_path)
table = doc.tables[1]
updates = {
    "Frozen 13-gene state vs OS": ("1.006", "0.874", "1.159", "0.909", "Null; modified HK"),
    "3-gene axis vs neutrophil module": ("0.400", "0.230", "0.546", "0.00344", "Replicated; modified HK"),
    "PLEK vs neutrophil module": ("0.693", "0.346", "0.873", "0.00857", "Primary driver; modified HK"),
    "PPBP vs neutrophil module": ("0.290", "0.209", "0.367", "0.000653", "Stable secondary bulk component; modified HK"),
    "PF4 vs neutrophil module": ("-0.00371", "-0.0896", "0.0822", "0.911", "Null component; modified HK"),
}
for row in table.rows[1:]:
    key = row.cells[0].text.strip()
    if key == "PLEK vs neutrophil module" and row.cells[1].text.strip() == "Spatial":
        vals = [key, "Spatial", "rho", "0.0806", "0.0410", "0.1199", "0.000135", "Weak; high heterogeneity; not specific"]
        for cell, value in zip(row.cells, vals):
            cell.text = value
        continue
    if key in updates:
        est, lo, hi, p, interpretation = updates[key]
        vals = [key, row.cells[1].text, row.cells[2].text, est, lo, hi, p, interpretation]
        for cell, value in zip(row.cells, vals):
            cell.text = value
for paragraph in doc.paragraphs:
    if paragraph.text.strip() == "Table 2. Key quantitative findings.":
        paragraph.text = ("Table 2. Key quantitative findings. Five-cohort meta-analytic estimates use REML "
                          "with modified Hartung–Knapp variance protection, t(k−1) inference, and prediction intervals; "
                          "the displayed P values and confidence intervals are the corrected values.")
doc.save(doc_path)
print(doc_path)
