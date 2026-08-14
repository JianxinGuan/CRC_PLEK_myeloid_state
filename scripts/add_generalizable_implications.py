from docx import Document
from pathlib import Path

pkg = Path(r"G:\New_CRC_Platelet\24_制作投稿文件包\manuscript\CRC_PLEK_myeloid_state_Main_Manuscript.docx")
doc = Document(pkg)

abstract_old = "The composite state is therefore better interpreted as a PLEK-associated myeloid context with neutrophil-related enrichment than as platelet abundance, a prognostic biomarker, or evidence of a causal platelet-neutrophil mechanism."
abstract_new = "More generally, composite bulk transcriptional states require component-level source auditing before their labels are interpreted as cell abundance or biological identity."

conclusion_new = (
    "In conclusion, this study does not identify a new function or cellular specificity for PLEK. "
    "It shows that a nominally platelet-related CRC tissue signature is substantially reclassified by a PLEK-rich myeloid compositional component with neutrophil-related enrichment. "
    "PPBP is a weaker, source-ambiguous bulk component, and PF4 contributes no stable independent association. "
    "The composite state is neither prognostic nor platelet-specific, and its spatial support is limited. "
    "The generalizable lesson is that composite transcriptional states require component-level source dissection before their names are allowed to imply their biology. "
    "In particular, bulk platelet-related scores should not be interpreted as platelet abundance unless their cellular source has been audited."
)

implications = (
    "Beyond this specific state, the findings provide a practical caution for bulk tumor transcriptomics. "
    "Platelet-, neutrophil-, and other immune-related gene sets are often scored in bulk tissue and interpreted as proxies for the corresponding cell populations. "
    "In the present analysis, a nominally platelet-related score carried its reproducible association largely through a broadly myeloid gene, whereas its platelet-identity/adhesion component did not show a stable cross-cohort association. "
    "Bulk cell-related scores should therefore be accompanied by a single-cell or orthogonal source audit, or explicitly described as having an unverified cellular origin, before they are interpreted as measures of cell abundance or cell-type-specific biology. "
    "Supplementary Table S1 operationalizes this interpretive discipline by listing the specific claims that are not supported by the present evidence. "
    "Together with the main-text interpretation boundaries in Table 3, it provides a transparent record that may be adapted when reporting other bulk-derived cellular association signals."
)

for p in doc.paragraphs:
    if abstract_old in p.text:
        p.text = p.text.replace(abstract_old, abstract_new)
    if p.text.startswith("In conclusion, this study does not identify a new function"):
        p.text = conclusion_new

# Insert the implications paragraph immediately before the limitations paragraph.
for p in list(doc.paragraphs):
    if p.text.startswith("The study remains observational."):
        previous = p._element.getprevious()
        if previous is None or "Beyond this specific state" not in "".join(previous.itertext()):
            new_p = p._parent.add_paragraph()
            p._element.addprevious(new_p._element)
            new_p.add_run(implications)
        break

doc.save(pkg)
print(pkg)
