from pathlib import Path
import re
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

root = Path(r"G:\New_CRC_Platelet")
manuscript = root / "24_制作投稿文件包" / "manuscript" / "CRC_PLEK_myeloid_state_Main_Manuscript.docx"
doc = Document(manuscript)

old = "PLEK showed the largest association, PPBP was secondary, and PF4 was null."
new = "PLEK, broadly expressed across myeloid cells, showed the strongest association; PPBP was weaker and PF4 null."
spatial_sentence = ("The spatial layer therefore acts as a constraint rather than a confirmation: it did not support a strong or consistently organized platelet-neutrophil colocalization signal, and it cannot establish direct contact or mechanism.")
for p in doc.paragraphs:
    p.text = p.text.replace("Among the three genes, the association was strongest for PLEK, which was broadly expressed across myeloid cells; PPBP showed a smaller bulk association, whereas PF4 was null.", new)
    if old in p.text:
        p.text = p.text.replace(old, new)
    if p.text.startswith("Spatial analysis placed a further limit on interpretation.") and spatial_sentence not in p.text:
        p.text = p.text.rstrip() + " " + spatial_sentence
doc.save(manuscript)

donor = pd.read_csv(root / "18_空间分析" / "results" / "06_donor_combined_spatial_effects.csv")
meta = pd.read_csv(root / "18_空间分析" / "results" / "07_donor_random_effects_spatial_meta.csv")
targets = [("PLEK", "PLEK-neutrophil"), ("granule3_axis", "Three-gene axis-neutrophil")]
fig, axes = plt.subplots(1, 2, figsize=(10.5, 5.4), sharey=True, constrained_layout=True)
plt.rcParams.update({"font.family":"Arial", "font.size":9})
for ax, (target, title) in zip(axes, targets):
    d = donor[(donor.target == target) & (donor.source == "neutrophil_module")].copy()
    d = d.sort_values("donor_id", ascending=False)
    d["lo"] = np.tanh(d.fisher_z - 1.96*d.se_z)
    d["hi"] = np.tanh(d.fisher_z + 1.96*d.se_z)
    y = np.arange(len(d)) + 1
    ax.errorbar(d.rho, y, xerr=np.vstack([d.rho-d.lo, d.hi-d.rho]), fmt='o', color="#366f91", ecolor="#7894a6", capsize=3, lw=1.2, ms=5)
    m = meta[(meta.target == target) & (meta.source == "neutrophil_module") & (meta.effect == "same_spot")].iloc[0]
    ax.errorbar(m.pooled_rho, 0, xerr=[[m.pooled_rho-m.rho_ci_low],[m.rho_ci_high-m.pooled_rho]], fmt='D', color="#b84a3a", capsize=4, lw=1.8, ms=6)
    ax.axvline(0, color="#555555", ls="--", lw=1)
    ax.set_yticks(np.r_[0,y], ["REML pooled"] + d.donor_id.tolist())
    ax.set_xlabel("Same-spot Spearman rho")
    ax.set_title(title, fontweight="bold")
    ax.grid(axis="x", color="#dddddd", lw=.7)
    ax.spines[['top','right','left']].set_visible(False)
fig.suptitle("Patient-level spatial-context associations", fontsize=13, fontweight="bold")

pkg_sup = root / "24_制作投稿文件包" / "supplementary"
repo_fig = root / "23_形成论文仓库" / "figures"
for folder in [pkg_sup, repo_fig]:
    folder.mkdir(exist_ok=True)
    fig.savefig(folder / "Supplementary_Figure_S3.pdf", bbox_inches="tight")
    fig.savefig(folder / "Supplementary_Figure_S3.png", dpi=300, bbox_inches="tight", facecolor="white")
plt.close(fig)

# Replace the embedded S3 and clarify its uncertainty display.
doc = Document(manuscript)
for i, p in enumerate(doc.paragraphs):
    if p.text.startswith("Supplementary Figure S3."):
        body = ("Patient-level random-effects spatial-context assessment. Same-spot spatial associations were small, "
                "heterogeneous, and not specific to the myeloid/neutrophil-related marker module. Patient-level error "
                "bars are approximate 95% confidence intervals derived on the Fisher-z scale; diamonds show REML-pooled estimates.")
        p.text = ""
        r1 = p.add_run("Supplementary Figure S3."); r1.bold = True
        r2 = p.add_run(" " + body); r2.bold = False
        image_p = doc.paragraphs[i-1]
        for run in list(image_p.runs):
            for drawing in list(run._element.xpath('.//w:drawing')):
                drawing.getparent().remove(drawing)
        image_p.add_run().add_picture(str(pkg_sup / "Supplementary_Figure_S3.png"), width=Inches(6.7))
        break
doc.save(manuscript)

# Copy exact plotted source data into repository figure_data.
donor_out = pd.concat([donor[(donor.target == t) & (donor.source == "neutrophil_module")].assign(plot_label=lab) for t,lab in targets])
donor_out.to_csv(root / "23_形成论文仓库" / "figure_data" / "Supplementary_Figure_S3_patient_data.csv", index=False)
meta_out = meta[(meta.target.isin([x[0] for x in targets])) & (meta.source == "neutrophil_module") & (meta.effect == "same_spot")]
meta_out.to_csv(root / "23_形成论文仓库" / "figure_data" / "Supplementary_Figure_S3_meta_data.csv", index=False)
print(manuscript)
