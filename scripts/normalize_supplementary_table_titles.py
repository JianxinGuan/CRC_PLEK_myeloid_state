from docx import Document
from pathlib import Path

p = Path(r"G:\New_CRC_Platelet\23_形成论文仓库\manuscript\CRC_PLEK_myeloid_state_manuscript_major_revision.docx")
d = Document(p)
titles = {
    1: "Claims prohibited by the staged evidence boundary.",
    2: "Cohort-level PLEK correlations with the eight-gene myeloid/neutrophil-related marker module and the matched seven-gene module excluding MNDA. FDR values use Benjamini-Hochberg correction across the ten cohort-version tests.",
    3: "REML meta-analysis with modified Hartung-Knapp variance protection, t(k-1) inference, and prediction intervals.",
    4: "Audit trail for the version-controlled staged analysis locks. The locked-file column uses repository-relative paths so that entries remain portable across systems. SHA-256 values identify the exact file contents at the recorded local timestamp. These internal locks document staged analysis decisions and do not constitute external prospective preregistration.",
    5: "Donor-level single-cell associations accounting for study-level dependence.",
    6: "Leave-one-study-out summary across the 29 study sources represented in the tumor donor pseudobulk analysis.",
    7: "GSE41258 coverage audit for the 13-gene platelet-related state.",
    8: "Modified Hartung-Knapp survival meta-analysis after excluding GSE41258.",
}
paras = d.paragraphs
for i in range(1, 9):
    standalone = f"Supplementary Table S{i}"
    for p0 in list(paras):
        if p0.text.strip() == standalone:
            el = p0._element
            el.getparent().remove(el)
    matches = [p0 for p0 in paras if p0.text.strip().startswith(f"Supplementary Table S{i}.")]
    if not matches:
        continue
    p0 = matches[0]
    p0.text = ""
    r1 = p0.add_run(f"Supplementary Table S{i}.")
    r1.bold = True
    r2 = p0.add_run(f" {titles[i]}")
    r2.bold = False
d.save(p)
print(p)
