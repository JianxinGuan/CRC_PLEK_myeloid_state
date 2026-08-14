from docx import Document
from pathlib import Path

path = Path(r"G:\New_CRC_Platelet\23_形成论文仓库\manuscript\CRC_PLEK_myeloid_state_Main_Manuscript.docx")
doc = Document(path)
for p in doc.paragraphs:
    if p.text.startswith("Supplementary Figure S3."):
        body = p.text[len("Supplementary Figure S3."):].strip()
        p.text = ""
        r1 = p.add_run("Supplementary Figure S3.")
        r1.bold = True
        r2 = p.add_run(" " + body)
        r2.bold = False
doc.save(path)
print(path)
