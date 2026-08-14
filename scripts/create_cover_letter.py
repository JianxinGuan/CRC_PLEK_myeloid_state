from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r"G:\New_CRC_Platelet\24_制作投稿文件包\Cover_Letter.docx")
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1); sec.bottom_margin = Inches(1); sec.left_margin = Inches(1); sec.right_margin = Inches(1)
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
style.font.size = Pt(12)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; p.add_run("14 August 2026")
doc.add_paragraph("Dear Editor,")
doc.add_paragraph("We submit the manuscript entitled “Deconstructing a platelet-related composite transcriptional state in colorectal cancer: cross-cohort evidence for a PLEK-associated myeloid tissue context” for consideration as an Article in Scientific Reports.")
doc.add_paragraph("Platelet-related transcripts in bulk colorectal cancer tissue are difficult to interpret because they may reflect platelets, residual blood, ambient RNA, or endogenous expression by non-platelet cells. In this study, we therefore used a deconstructive biological-association design rather than treating a platelet-related gene score as a direct measure of platelet abundance. A frozen 13-gene state was separated into prespecified components and evaluated across five GEO cohorts, TCGA-COAD/READ, a large CRC single-cell atlas, and a spatial transcriptomic dataset.")
doc.add_paragraph("The main result is a constrained one. The 13-gene state was not associated with overall survival. The reproducible association was concentrated in the PF4/PPBP/PLEK component, with unequal gene contributions: PLEK showed the largest association with a neutrophil-related bulk module, PPBP showed a smaller bulk signal, and PF4 was null. Purity-adjusted TCGA analysis, single-cell source auditing, donor-level pseudobulk analyses, and study-level sensitivity analyses supported a broad PLEK-associated myeloid context with neutrophil-related enrichment. Sparse PPBP/PF4 detection in neutrophils and weak, heterogeneous spatial associations prevented interpretation as platelet-specific expression, a neutrophil-specific source, direct platelet-neutrophil contact, a prognostic biomarker, or a causal mechanism.")
doc.add_paragraph("We believe the manuscript is suitable for Scientific Reports because it addresses a common interpretive problem in bulk tumor transcriptomics: how an apparently platelet-related signal changes when its cellular context and component contributions are examined across independent data types. The study preserves negative findings, reports heterogeneity and prediction intervals, and explicitly distinguishes reproducible association from biological source or mechanism.")
doc.add_paragraph("This manuscript has not been published and is not under consideration elsewhere. The work reanalyzes de-identified data from public repositories and did not involve new human recruitment or tissue collection. The authors declare no competing interests and received no specific grant funding. All authors have reviewed and approved the submitted manuscript. The corresponding author is Yongbin Qin (yongbinqin1988@163.com), Yulin Red Cross Hospital, Yulin, Guangxi, China.")
doc.add_paragraph("The analysis code, frozen analysis specifications, derived figure data, and audit records are available in the accompanying repository: https://github.com/JianxinGuan/CRC_PLEK_myeloid_state.")
doc.add_paragraph("Thank you for considering our manuscript.")
doc.add_paragraph("Sincerely,")
doc.add_paragraph("Yongbin Qin")
doc.add_paragraph("Corresponding author")
doc.add_paragraph("Yulin Red Cross Hospital")
doc.add_paragraph("Yulin, Guangxi, China")
doc.add_paragraph("Email: yongbinqin1988@163.com")
doc.save(out)
print(out)
