from docx import Document
from pathlib import Path

path = Path(r"G:\New_CRC_Platelet\24_制作投稿文件包\manuscript\CRC_PLEK_myeloid_state_Main_Manuscript.docx")
doc = Document(path)

# The first three document tables are main Tables 1-3. Remove embedded S1-S8 table bodies.
for table in list(doc.tables)[3:]:
    element = table._element
    element.getparent().remove(element)

s4 = (
    "Audit trail for the version-controlled staged analysis locks. Repository-relative paths "
    "and SHA-256 values identify the locked files. These internal records do not constitute "
    "prospective external preregistration."
)
for paragraph in doc.paragraphs:
    if paragraph.text.startswith("Supplementary Table S4."):
        paragraph.text = ""
        label = paragraph.add_run("Supplementary Table S4.")
        label.bold = True
        body = paragraph.add_run(" " + s4)
        body.bold = False

doc.save(path)
print(path)
print(f"remaining tables: {len(doc.tables)}")
