from docx import Document
from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

path = Path(r"G:\New_CRC_Platelet\24_制作投稿文件包\manuscript\CRC_PLEK_myeloid_state_Main_Manuscript.docx")
doc = Document(path)

# Rebuild Table 2 with the estimate and CI combined, reducing eight columns to six.
old = doc.tables[1]
rows = []
for row in old.rows[1:]:
    c = [x.text.strip() for x in row.cells]
    rows.append([c[0], c[1], c[2], f"{c[3]} ({c[4]} to {c[5]})", c[6], c[7]])
new = doc.add_table(rows=1, cols=6)
new.style = "Table Grid"
headers = ["Comparison", "Evidence layer", "Effect measure", "Estimate (95% CI)", "P value/FDR", "Interpretation"]
for cell, value in zip(new.rows[0].cells, headers): cell.text = value
for vals in rows:
    cells = new.add_row().cells
    for cell, value in zip(cells, vals): cell.text = value
old._element.addprevious(new._element)
old._element.getparent().remove(old._element)

# Human-readable headers for Tables 1 and 3.
header_maps = {
    0: ["Analysis layer", "Platform", "Independent units", "Analysis observations", "Role"],
    2: ["Evidence status", "Permitted interpretation"],
}
for idx, vals in header_maps.items():
    for cell, value in zip(doc.tables[idx].rows[0].cells, vals): cell.text = value

for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        trPr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)
        if r_idx == 0:
            tbl_header = OxmlElement('w:tblHeader')
            tbl_header.set(qn('w:val'), 'true')
            trPr.append(tbl_header)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), 'D9E1F2'); cell._tc.get_or_add_tcPr().append(shd)
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if r_idx == 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)
                    run.bold = (r_idx == 0)

doc.save(path)
print(path)
