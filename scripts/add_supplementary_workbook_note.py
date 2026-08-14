from docx import Document
from pathlib import Path

path = Path(r"G:\New_CRC_Platelet\24_制作投稿文件包\manuscript\CRC_PLEK_myeloid_state_Main_Manuscript.docx")
doc = Document(path)
note = "Supplementary Tables S1–S8 are provided in the accompanying Excel workbook."
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "Supplementary Tables":
        if i + 1 < len(doc.paragraphs) and doc.paragraphs[i + 1].text.strip() != note:
            new_p = p._parent.add_paragraph()
            p._element.addnext(new_p._element)
            new_p.add_run(note)
        break
doc.save(path)
print(path)
