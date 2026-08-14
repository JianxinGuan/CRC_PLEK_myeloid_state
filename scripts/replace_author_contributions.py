from docx import Document
from pathlib import Path

root = Path(r"G:\New_CRC_Platelet\23_形成论文仓库\manuscript")
new_text = (
    "J.G. and S.J.: conceptualization, methodology, software, formal analysis, data curation, "
    "visualization, and writing—original draft. J.C. and Y.D.: data curation, validation, "
    "investigation, and review and editing. Y.Q.: supervision, project administration, "
    "resources, and review and editing. All authors reviewed and approved the final manuscript."
)
for path in root.glob("*.docx"):
    doc = Document(path)
    replaced = False
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Author contributions":
            if i + 1 < len(doc.paragraphs) and ("AUTHOR CONTRIBUTIONS REQUIRE" in doc.paragraphs[i + 1].text or "contributed equally" in doc.paragraphs[i + 1].text):
                doc.paragraphs[i + 1].text = new_text
                replaced = True
    if replaced:
        doc.save(path)
        print(path)
